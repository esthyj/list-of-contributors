import json
from docx import Document
from docx.shared import Pt, RGBColor, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx2pdf import convert

txt_filename = "my_data.txt"
word_filename = "report.docx"
pdf_filename = "report.pdf"

try:
    with open(txt_filename, "r", encoding="utf-8") as f:
        data_dict = json.load(f)
    print(f"'{txt_filename}'에서 데이터를 성공적으로 불러왔습니다.")

    date = data_dict.get("날짜", "날짜 없음")
    if "날짜" in data_dict:
        del data_dict["날짜"]

    # 새로운 문서
    document = Document()

    # A3 가로 및 여백 설정
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Mm(420.0)
    section.page_height = Mm(297.0)
    section.top_margin = Mm(15.0)
    section.header_distance = Mm(5.0)
    section.bottom_margin = Mm(5.0)
    section.footer_distance = Mm(5.0)
    section.left_margin = Mm(25.0)
    section.right_margin = Mm(25.0)
    section.gutter = Mm(0.0)

    # 1. 큰 제목: "헌금자 명단" (40pt, Bold)
    title_paragraph = document.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_paragraph.paragraph_format.space_after = Pt(0)
    run = title_paragraph.add_run("헌금자 명단")
    run.font.size = Pt(40)
    run.font.bold = True
    run.font.name = "함초롬바탕"
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:eastAsia"), "함초롬바탕")

    # 2. 작은 제목 (날짜) (20pt, 일반)
    subtitle_paragraph = document.add_paragraph()
    subtitle_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle_paragraph.add_run("(" + str(date) + ")")
    sub_run.font.size = Pt(25)
    sub_run.font.bold = True
    sub_run.font.name = "함초롬바탕"
    sub_rPr = sub_run._r.get_or_add_rPr()
    sub_rFonts = sub_rPr.get_or_add_rFonts()
    sub_rFonts.set(qn("w:eastAsia"), "함초롬바탕")
    document.add_paragraph()

    # 3. 표 생성
    NUM_COLUMNS = 6
    col_width = Mm(61.5)  # 균등 열 너비

    for key, names in data_dict.items():
        key_paragraph = document.add_paragraph()
        key_run = key_paragraph.add_run(key)
        key_run.font.size = Pt(30)
        key_run.font.bold = True
        key_run.font.color.rgb = RGBColor(0, 0, 0)
        key_run.font.name = "함초롬바탕"
        key_rPr = key_run._r.get_or_add_rPr()
        key_rFonts = key_rPr.get_or_add_rFonts()
        key_rFonts.set(qn("w:eastAsia"), "함초롬바탕")
        key_format = key_paragraph.paragraph_format
        key_format.space_after = Pt(0)

        if not names:
            document.add_paragraph("(내역 없음)")
            document.add_paragraph()
            continue

        # 표 생성
        num_rows = -(-len(names) // NUM_COLUMNS)
        table = document.add_table(rows=num_rows, cols=NUM_COLUMNS)

        # ✅ autofit 비활성화 (열 너비 고정)
        table.autofit = False
        table.allow_autofit = False

        # ✅ XML로 테이블 레이아웃을 'fixed'로 설정
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
        tblLayout = OxmlElement("w:tblLayout")
        tblLayout.set(qn("w:type"), "fixed")
        tblPr.append(tblLayout)

        # ✅ 열 너비 설정
        for col in table.columns:
            col.width = col_width

        # 표 테두리 투명으로 변경
        tblBorders = OxmlElement("w:tblBorders")
        for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
            border_el = OxmlElement(f"w:{border_name}")
            border_el.set(qn("w:val"), "none")
            border_el.set(qn("w:sz"), "0")
            tblBorders.append(border_el)
        tblPr.append(tblBorders)

        # 표 채우기
        name_index = 0
        for i in range(num_rows):
            for j in range(NUM_COLUMNS):
                cell = table.cell(i, j)

                # ✅ 각 셀의 너비도 명시적으로 설정
                cell.width = col_width

                if name_index < len(names):
                    cell.text = names[name_index]

                    paragraph = cell.paragraphs[0]
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    p_format = paragraph.paragraph_format
                    p_format.space_before = Pt(0)
                    p_format.space_after = Pt(0)

                    cell_run = cell.paragraphs[0].runs[0]
                    cell_run.font.size = Pt(25)
                    cell_run.font.bold = False
                    cell_run.font.name = "함초롬바탕"
                    cell_rPr = cell_run._r.get_or_add_rPr()
                    cell_rFonts = cell_rPr.get_or_add_rFonts()
                    cell_rFonts.set(qn("w:eastAsia"), "함초롬바탕")

                    name_index += 1
                else:
                    cell.text = ""
        document.add_paragraph()

    # 4. 로고 추가
    print("문서 마지막에 이미지를 추가합니다...")
    try:
        document.add_picture("logo.png", width=Mm(130))
        last_paragraph = document.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except FileNotFoundError:
        print("🚨 경고: 'logo.png' 파일을 찾지 못해 이미지 추가를 건너뜁니다.")
    except Exception as e:
        print(f"🚨 경고: 이미지 추가 중 오류 발생: {e}")

    # 5. 파일 저장
    document.save(word_filename)
    print(f"\n✅ 두 줄 제목이 적용된 '{word_filename}' 파일이 저장되었습니다.")
    print(f"\nPDF 변환을 시작합니다 (Word/LibreOffice 필요)...")
    try:
        # 현재 폴더의 .docx 파일을 .pdf 파일로 변환
        convert(word_filename, pdf_filename)
        print(f"✅ PDF 파일 '{pdf_filename}' 저장이 완료되었습니다.")

    except Exception as e:
        print(f"\n🚨 PDF 변환 중 오류가 발생했습니다: {e}")
        print(
            "  [문제 해결] PC에 Microsoft Word 또는 LibreOffice가 설치되어 있는지 확인하세요."
        )
        print("  (Windows의 경우 Word, macOS/Linux의 경우 LibreOffice가 필요합니다.)")

except FileNotFoundError:
    print(f"🚨 오류: '{txt_filename}' 파일을 찾을 수 없습니다.")
except Exception as e:
    print(f"🚨 오류 발생: {e}")
