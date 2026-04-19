"""헌금자 명단 처리: RTF 파싱 → Word 문서(.doc) 생성."""

import os
import re
from striprtf.striprtf import rtf_to_text
from docx import Document
from docx.shared import Pt, RGBColor, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.section import WD_ORIENT
import win32com.client

# ── 설정 ──────────────────────────────────────────────
DATA_DIR = "data"
FOLDER_PATH = DATA_DIR
LOGO_FILE = os.path.join(DATA_DIR, "logo.png")
FONT_NAME = "함초롬바탕"
NUM_COLUMNS = 6
COL_WIDTH = Mm(61.5)

EXCLUDED_KEYWORDS = ["주일헌금", "특별헌금", "공과금", "구제헌금", "헌물"]
PRIORITY_KEYS = ["날짜", "십일조", "감사헌금", "일천번제"]


# ── 1. RTF 파싱 ──────────────────────────────────────

def decode_bytes_cp949(b: bytes) -> str:
    """cp949 우선으로 바이트를 문자열로 디코딩한다."""
    for enc in ("cp949", "utf-8", "euc-kr", "latin-1"):
        try:
            return b.decode(enc)
        except Exception:
            continue
    return b.decode("cp949", errors="replace")


def find_rtf_files(folder: str) -> list[str]:
    """폴더에서 .rtf 파일 목록을 반환한다."""
    return [f for f in os.listdir(folder) if f.lower().endswith(".rtf")]


def extract_text_from_rtf(folder: str, filenames: list[str]) -> str:
    """RTF 파일들에서 텍스트를 추출하여 합친다."""
    text = ""
    for filename in filenames:
        full_path = os.path.join(folder, filename)
        with open(full_path, "rb") as f:
            raw = f.read()
        rtf_str = decode_bytes_cp949(raw)
        text += rtf_to_text(rtf_str)
    return text


def extract_date(text: str) -> str:
    """텍스트에서 날짜(YYYY.M.D) 패턴을 찾아 반환한다."""
    match = re.search(r"\d{4}\.\d{1,2}\.\d{1,2}", text)
    return match.group(0) if match else "날짜 없음"


def pad_date(date: str) -> str:
    """날짜를 'YYYY.MM.DD' 형식(제로 패딩)으로 정규화한다."""
    parts = date.split(".")
    if len(parts) == 3 and all(p.strip().isdigit() for p in parts):
        return f"{int(parts[0]):04d}.{int(parts[1]):02d}.{int(parts[2]):02d}"
    return date


def extract_offering_section(text: str) -> str:
    """구분선(-----)으로 나뉜 텍스트에서 헌금 내역 부분만 추출한다."""
    parts = re.split(r"-{5,}", text)
    if len(parts) == 5:
        return parts[1].strip() + "\n" + parts[3].strip()
    elif len(parts) == 3:
        return parts[1].strip()
    else:
        print(f"경고: 구분선 기준으로 {len(parts)}개 부분으로 나뉘었습니다.")
        return ""


def parse_offerings(text: str) -> dict[str, list[str]]:
    """헌금 내역 텍스트를 파싱하여 {항목: [이름들]} 딕셔너리로 반환한다."""
    result = {}
    current_key = None
    parts = re.split(r"\s*(\S+\s\*\s\d+\s명)\s*", text.strip())

    for part in parts:
        if not part:
            continue
        if "*" in part and "명" in part:
            current_key = part.split("*")[0].strip()
            result[current_key] = []
        elif current_key:
            names = re.split(r"\s{2,}", part.strip())
            result[current_key].extend(names)

    return result


def fix_combined_keys(d: dict) -> dict:
    """줄바꿈 등으로 합쳐진 키 이름을 정리한다. (예: '감사헌금십일조' → '십일조')"""
    pattern = re.compile(r"(.*?)(헌금|번제|첫열매|건물리모델링|설립)")
    renamed = {}

    for old_key, value in d.items():
        matches = pattern.findall(old_key)
        if len(matches) > 1:
            new_key = "".join(matches[-1])
            if new_key in renamed:
                renamed[new_key].extend(value)
            else:
                renamed[new_key] = value
        else:
            if old_key in renamed:
                renamed[old_key].extend(value)
            else:
                renamed[old_key] = value

    return renamed


def remove_excluded(d: dict) -> dict:
    """제외 대상 키워드가 포함된 키/값을 제거한다."""
    result = {}
    for key, names in d.items():
        if any(kw in key for kw in EXCLUDED_KEYWORDS):
            continue
        filtered = [n for n in names if not any(kw in n for kw in EXCLUDED_KEYWORDS)]
        result[key] = filtered
    return result


def format_result(d: dict, date: str) -> dict:
    """키에 인원수를 추가하고, 날짜를 넣고, 우선순위대로 정렬한다."""
    keyed = {f"{k} ({len(v)}명)": v for k, v in d.items()}
    keyed["날짜"] = date

    ordered = {}
    for priority in PRIORITY_KEYS:
        for k, v in keyed.items():
            if priority in k and k not in ordered:
                ordered[k] = v
    for k, v in keyed.items():
        if k not in ordered:
            ordered[k] = v

    return ordered


def parse_rtf() -> dict:
    """RTF 파일을 파싱하여 JSON 파일로 저장하고 결과를 반환한다."""
    rtf_files = find_rtf_files(FOLDER_PATH)
    if not rtf_files:
        print("RTF 파일이 없습니다.")
        return {}

    print(f"RTF 파일 {len(rtf_files)}개 발견: {rtf_files}")
    text = extract_text_from_rtf(FOLDER_PATH, rtf_files)

    date = extract_date(text)
    section = extract_offering_section(text)
    section = section.replace("첫 열 매", "첫열매")

    offerings = parse_offerings(section)
    offerings = fix_combined_keys(offerings)
    offerings = remove_excluded(offerings)
    return format_result(offerings, date)


# ── 2. Word 문서 생성 ────────────────────────────────

def set_font(run, size: int, bold: bool = False):
    """run에 폰트를 설정한다."""
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = FONT_NAME
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:eastAsia"), FONT_NAME)


def setup_page(document: Document):
    """A3 가로 페이지 설정."""
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Mm(420.0)
    section.page_height = Mm(297.0)
    section.top_margin = Mm(15.0)
    section.header_distance = Mm(5.0)
    section.bottom_margin = Mm(5.0)
    section.footer_distance = Mm(5.0)
    section.left_margin = Mm(25.0)
    section.right_margin = Mm(25.0)
    section.gutter = Mm(0.0)


def add_title(document: Document, date: str):
    """제목과 날짜를 추가한다."""
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(0)
    set_font(title.add_run("헌금자 명단"), size=40, bold=True)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(subtitle.add_run(f"({date})"), size=25, bold=True)

    document.add_paragraph()


def make_table_borderless(table):
    """테이블 테두리를 투명하게 설정한다."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")

    tblLayout = OxmlElement("w:tblLayout")
    tblLayout.set(qn("w:type"), "fixed")
    tblPr.append(tblLayout)

    tblBorders = OxmlElement("w:tblBorders")
    for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        el = OxmlElement(f"w:{border_name}")
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        tblBorders.append(el)
    tblPr.append(tblBorders)


def add_offering_section(document: Document, title: str, names: list[str]):
    """헌금 항목 하나를 제목 + 표로 추가한다."""
    heading = document.add_paragraph()
    run = heading.add_run(title)
    run.font.color.rgb = RGBColor(0, 0, 0)
    set_font(run, size=30, bold=True)
    heading.paragraph_format.space_after = Pt(0)

    if not names:
        document.add_paragraph("(내역 없음)")
        document.add_paragraph()
        return

    num_rows = -(-len(names) // NUM_COLUMNS)
    table = document.add_table(rows=num_rows, cols=NUM_COLUMNS)
    table.autofit = False
    table.allow_autofit = False

    for col in table.columns:
        col.width = COL_WIDTH

    make_table_borderless(table)

    for i, name in enumerate(names):
        row, col = divmod(i, NUM_COLUMNS)
        cell = table.cell(row, col)
        cell.width = COL_WIDTH
        cell.text = name

        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)

        set_font(paragraph.runs[0], size=25)

    document.add_paragraph()


def add_logo(document: Document):
    """문서 끝에 로고 이미지를 추가한다."""
    try:
        document.add_picture(LOGO_FILE, width=Mm(130))
        document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    except FileNotFoundError:
        print(f"경고: '{LOGO_FILE}' 파일을 찾지 못해 이미지를 건너뜁니다.")


def create_docx(data: dict) -> tuple[str, str]:
    """딕셔너리 데이터로 .docx 문서를 생성한다. (docx 경로, doc 경로) 반환."""
    date = data.pop("날짜", "날짜 없음")
    padded = pad_date(date)
    base_name = f"헌금자명단 ({padded})"
    docx_file = f"{base_name}.docx"
    doc_file = f"{base_name}.doc"

    document = Document()
    setup_page(document)
    add_title(document, date)

    for key, names in data.items():
        add_offering_section(document, key, names)

    add_logo(document)
    document.save(docx_file)
    print(f"'{docx_file}' 저장 완료.")

    return docx_file, doc_file


def convert_docx_to_doc(docx_file: str, doc_file: str):
    """.docx 파일을 .doc 형식으로 변환한다 (Microsoft Word 필요)."""
    docx_path = os.path.abspath(docx_file)
    doc_path = os.path.abspath(doc_file)

    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False
    try:
        doc = word.Documents.Open(docx_path)
        doc.SaveAs(doc_path, FileFormat=0)  # 0 = wdFormatDocument (.doc)
        doc.Close()
        print(f"'{doc_file}' 변환 완료.")
    finally:
        word.Quit()

    os.remove(docx_file)


# ── 메인 실행 ─────────────────────────────────────────

def main():
    print("▶ RTF 파일 파싱 중…")
    data = parse_rtf()
    if not data:
        return
    print("✓ 파싱 완료!\n")

    print("▶ Word 문서 생성 중…")
    docx_file, doc_file = create_docx(data)
    convert_docx_to_doc(docx_file, doc_file)
    print("✓ 문서 생성 완료!")

    print("\n✅ 모든 작업 완료!")


if __name__ == "__main__":
    main()
